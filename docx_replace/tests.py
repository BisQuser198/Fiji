from django.test import TestCase

# Create your tests here.

# test extracting text from a docx file
import os
from docx import Document

# Objective: I want to use a hardcoded path to the project root, but it needs to work on any machine

# 1) Get the absolute path of the current file
current_file = os.path.abspath(__file__)
# 2) Get the directory containing the current file (e.g., docx_replace)
current_dir = os.path.dirname(current_file)
# 3) # Get the parent directory (your project root, Fiji)
project_root = os.path.dirname(current_dir)
# 4) Print the project root to verify
print("Project Root:", project_root)
# 5) Build the path to the Raw Data folder and your document
raw_data_folder = os.path.join(project_root, 'Raw Data') # all test files are in here
doc_filename = 'test 1.docx'
doc_path = os.path.join(raw_data_folder, doc_filename)
# 6) Check I've reached my document:
print(doc_path)
# 7) Open my document and print its text content
doc = Document(doc_path)  ## # doc = Document('example.docx')
for para in doc.paragraphs:
    # print(para.text)
    # 8) print each run (run = segment of text with the same style / formatting)
    for run in para.runs:
        print(run.text)
