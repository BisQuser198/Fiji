
# docx_replace/utils.py

import logging
logger = logging.getLogger(__name__)


import os
import string
import zipfile
import tempfile
from io import BytesIO
from datetime import datetime

import pandas as pd
from pandas import Timestamp
from docx import Document

def replace_placeholder_preserve_runs(para, placeholder, replacement):
    """
    Scan through runs, detect placeholder spanning multiple runs,
    and rewrite each run’s text segments so formatting remains intact.
    Returns True if at least one replacement was done.
    """
    # Build cumulative mapping of runs → text offsets
    cumulative = []
    full_text = ""
    for i, run in enumerate(para.runs):
        text = run.text or ""
        start = len(full_text)
        full_text += text
        end = len(full_text)
        cumulative.append((i, start, end))

    if placeholder not in full_text:
        return False

    # Create the new full text with replacements
    new_full = full_text.replace(placeholder, replacement)

    # Clear all runs
    for run in para.runs:
        run.text = ""

    # Redistribute new_full across runs by original lengths
    pos = 0
    for i, start, end in cumulative:
        length = end - start
        slice_text = new_full[pos : pos + length]
        para.runs[i].text = slice_text
        pos += length

    # Append any leftover text to the last run
    if pos < len(new_full):
        para.runs[-1].text += new_full[pos:]

    return True


def col_letter_to_index(letter: str) -> int:
    """
    Convert Excel column letters (A, B, …, Z, AA, AB, …) to zero-based index.
    """
    letter = letter.strip().upper()
    idx = 0
    for ch in letter:
        if ch in string.ascii_uppercase:
            idx = idx * 26 + (ord(ch) - ord('A') + 1)
        else:
            raise ValueError(f"Invalid column letter: {letter}")
    return idx - 1

def batch_find_replace(
    excel_path: str,
    header_rows: int,
    id_col_letter: str,
    filename_pattern: str,
    start_id: int,
    end_id: int,
    replacements: list,     # list of [find_text, col_letter]
    docx_zip_bytes: bytes   # bytes of the uploaded ZIP
):
    """
    1) Reads Excel file from excel_path (skipping header_rows).
    2) Builds a map of ID -> row data using id_col_letter.
    3) Unzips the incoming DOCX ZIP (docx_zip_bytes) into a temp folder.
    4) For each doc_id in [start_id..end_id]:
         - Opens the matching DOCX (filename_pattern.format(id=doc_id))
         - Finds/replaces each placeholder, with debug logging
         - Saves modified docs into an output temp folder
    5) Zips up the output folder and returns (zip_bytes, logs_list).
    """
    # check if def batch_find_replace(...) is running at all
    logger.debug("▶︎ Entered batch_find_replace() with %d rules", len(replacements))


    # ---- 1) Read Excel & map IDs to rows ----
    df = pd.read_excel(excel_path, header=None, skiprows=header_rows)
    id_idx = col_letter_to_index(id_col_letter)
    id_to_row = {}
    for _, row in df.iterrows():
        raw = row[id_idx]
        if pd.isna(raw):
            continue
        key = int(raw) if isinstance(raw, (int, float)) and float(raw).is_integer() else raw
        id_to_row[key] = row

    # ---- 2) Unzip DOCX files into temp_in ----
    tmp_in = tempfile.mkdtemp(prefix="docx_in_")
    tmp_out = tempfile.mkdtemp(prefix="docx_out_")
    with zipfile.ZipFile(BytesIO(docx_zip_bytes)) as zin:
        zin.extractall(tmp_in)
    # test to see which files are being extracted
    files = os.listdir(tmp_in)
    logger.debug("▶︎ Extracted DOCX files: %r", files)


    logs = []

    # ---- 3) Process each ID ----
    for doc_id in range(start_id, end_id + 1):
        # check if files are being processed
        logger.debug("→ Processing ID %s", doc_id)
        fname = filename_pattern.format(id=doc_id)
        src_path = os.path.join(tmp_in, fname)

        if not os.path.isfile(src_path):
            logs.append(f"[{doc_id}] File not found: {fname}")
            continue
        if doc_id not in id_to_row:
            logs.append(f"[{doc_id}] No Excel data; skipping {fname}")
            continue

        doc = Document(src_path)
        row = id_to_row[doc_id]
        changed = False

        # ---- 4) Apply each find/replace rule ----
        for find_txt, col_let in replacements:
            col_idx = col_letter_to_index(col_let)
            raw_val = row[col_idx]

            if isinstance(raw_val, (Timestamp, datetime)):
                repl = raw_val.strftime("%d.%m.%Y")
            elif pd.isna(raw_val):
                repl = ""
            else:
                repl = str(raw_val)

            # --- Paragraphs ---
            for para in doc.paragraphs:
                if find_txt in para.text:
                    #print(f"[DEBUG] Found '{find_txt}' in paragraph: '{para.text[:40]}…'")
                    logger.debug("Found %r in paragraph: %.40r…", find_txt, para.text)
                    # Try run-by-run replace
                    replaced_in_run = False
                    for ri, run in enumerate(para.runs):
                        if find_txt in run.text:
                            #print(f"[DEBUG] ‑- Replacing in run {ri}: '{run.text}'")
                            logger.debug("Found %r in paragraph: %.40r…", find_txt, para.text)
                            run.text = run.text.replace(find_txt, repl)
                            changed = True
                            replaced_in_run = True

                    # Fallback: coalesce paragraph
                    # if not replaced_in_run:
                        #print(f"[WARN] Placeholder spans runs — coalescing paragraph.")
                        #logger.warning("Placeholder spans runs — coalescing paragraph.")
                        # full = para.text.replace(find_txt, repl)
                        # for run in para.runs:
                        #     run.text = ""
                        # para.runs[0].text = full
                        # changed = True
                    # Fallback: preserve formatting across runs
                    if not replaced_in_run:
                        logger.warning("Spanning-run placeholder; invoking style-preserving replacer")
                        if replace_placeholder_preserve_runs(para, find_txt, repl):
                            changed = True


            # --- Tables ---
            for table in doc.tables:
                for row_cells in table.rows:
                    for cell in row_cells.cells:
                        for para in cell.paragraphs:
                            if find_txt in para.text:
                                #print(f"[DEBUG] Found '{find_txt}' in table cell: '{para.text[:40]}…'")
                                logger.debug("Found %r in paragraph: %.40r…", find_txt, para.text)
                                replaced_in_run = False
                                for ri, run in enumerate(para.runs):
                                    if find_txt in run.text:
                                        #print(f"[DEBUG] ‑- Replacing in run {ri}: '{run.text}'")
                                        logger.debug("Found %r in paragraph: %.40r…", find_txt, para.text)
                                        run.text = run.text.replace(find_txt, repl)
                                        changed = True
                                        replaced_in_run = True
                                # if not replaced_in_run:
                                #     #print(f"[WARN] Table placeholder spans runs — coalescing paragraph.")
                                #     logger.warning("Placeholder spans runs — coalescing paragraph.")
                                #     full = para.text.replace(find_txt, repl)
                                #     for run in para.runs:
                                #         run.text = ""
                                #     para.runs[0].text = full
                                #     changed = True
                                # Fallback: preserve formatting across runs
                                if not replaced_in_run:
                                    logger.warning("Spanning-run placeholder; invoking style-preserving replacer")
                                    if replace_placeholder_preserve_runs(para, find_txt, repl):
                                        changed = True


        # ---- 5) Save to output folder ----
        dst_path = os.path.join(tmp_out, fname)
        if changed:
            doc.save(dst_path)
            logs.append(f"[{doc_id}] Updated: {fname}")
        else:
            # No changes: copy original
            from shutil import copyfile
            copyfile(src_path, dst_path)
            logs.append(f"[{doc_id}] Unchanged (no placeholder found): {fname}")

    # ---- 6) Zip up modified docs ----
    out_zip = BytesIO()
    with zipfile.ZipFile(out_zip, "w", zipfile.ZIP_DEFLATED) as zout:
        for root, _, files in os.walk(tmp_out):
            for f in files:
                full = os.path.join(root, f)
                arc = os.path.relpath(full, tmp_out)
                zout.write(full, arc)
    out_zip.seek(0)

    return out_zip.getvalue(), logs


## Backup copy of code - works until the find-and-replace part (without the replace)
# # docx_replace/utils.py
# import os, string, zipfile
# import pandas as pd
# from io import BytesIO
# from datetime import datetime
# from pandas import Timestamp
# from docx import Document

# def col_letter_to_index(letter: str) -> int:
#     letter = letter.strip().upper()
#     idx = 0
#     for ch in letter:
#         idx = idx*26 + (ord(ch)-ord('A')+1)
#     return idx-1

# def batch_find_replace(
#     excel_path: str,
#     header_rows: int,
#     id_col_letter: str,
#     filename_pattern: str,
#     start_id: int,
#     end_id: int,
#     replacements: list,     # list of (find_txt, col_letter)
#     docx_zip_bytes: bytes   # uploaded ZIP file as bytes
# ) -> bytes:
#     """
#     - Reads Excel from excel_path.
#     - Unzips docx_zip_bytes into a temp dir.
#     - Applies find/replace on each doc ID in [start_id, end_id].
#     - Re-zips the modified docs and returns ZIP bytes.
#     """
#     # 1) Read Excel
#     df = pd.read_excel(excel_path, header=None, skiprows=header_rows)
#     id_idx = col_letter_to_index(id_col_letter)
#     id_to_row = {}
#     for _, row in df.iterrows():
#         raw = row[id_idx]
#         if pd.isna(raw): continue
#         key = int(raw) if isinstance(raw,(int,float)) and float(raw).is_integer() else raw
#         id_to_row[key] = row

#     # 2) Unzip docs to tmp folder
#     import tempfile
#     tmp_in = tempfile.mkdtemp()
#     tmp_out = tempfile.mkdtemp()
#     with zipfile.ZipFile(BytesIO(docx_zip_bytes)) as zin:
#         zin.extractall(tmp_in)

#     logs = []
#     # 3) Process each ID
#     for doc_id in range(start_id, end_id+1):
#         fname = filename_pattern.format(id=doc_id)
#         src_path = os.path.join(tmp_in, fname)
#         if not os.path.isfile(src_path):
#             logs.append(f"[{doc_id}] Missing file {fname}")
#             continue
#         if doc_id not in id_to_row:
#             logs.append(f"[{doc_id}] No Excel data; skipping.")
#             continue

#         doc = Document(src_path)
#         row = id_to_row[doc_id]
#         changed = False

#         for find_txt, col_let in replacements:
#             col_idx = col_letter_to_index(col_let)
#             raw_val = row[col_idx]
#             if isinstance(raw_val,(Timestamp,datetime)):
#                 repl = raw_val.strftime("%d.%m.%Y")
#             elif pd.isna(raw_val):
#                 repl = ""
#             else:
#                 repl = str(raw_val)

#             # replace in paras & tables
#             for para in doc.paragraphs:
#                 for run in para.runs:
#                     if find_txt in run.text:
#                         run.text = run.text.replace(find_txt,repl)
#                         changed = True
#             for table in doc.tables:
#                 for r in table.rows:
#                     for cell in r.cells:
#                         for p in cell.paragraphs:
#                             for run in p.runs:
#                                 if find_txt in run.text:
#                                     run.text = run.text.replace(find_txt,repl)
#                                     changed = True

#         # save to output tree
#         dst = os.path.join(tmp_out, fname)
#         if changed:
#             doc.save(dst)
#             logs.append(f"[{doc_id}] Updated {fname}")
#         else:
#             # copy original
#             from shutil import copyfile
#             copyfile(src_path, dst)
#             logs.append(f"[{doc_id}] Unchanged {fname}")

#     # 4) Re-zip outputs
#     memory_zip = BytesIO()
#     with zipfile.ZipFile(memory_zip, 'w') as zout:
#         for root, _, files in os.walk(tmp_out):
#             for f in files:
#                 full = os.path.join(root,f)
#                 arc = os.path.relpath(full, tmp_out)
#                 zout.write(full, arc)
#     memory_zip.seek(0)
#     return memory_zip.getvalue(), logs
